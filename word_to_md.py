import os
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def convert_word_to_markdown(word_file_path, output_md_path=None):
    """
    将Word文档转换为Markdown文件
    
    Args:
        word_file_path: Word文档路径
        output_md_path: 输出的Markdown文件路径，如果为None则自动生成
    """
    if not os.path.exists(word_file_path):
        print(f"错误：文件 {word_file_path} 不存在")
        return
    
    # 如果没有指定输出路径，则自动生成
    if output_md_path is None:
        base_name = os.path.splitext(word_file_path)[0]
        output_md_path = base_name + '.md'
    
    # 读取Word文档
    doc = Document(word_file_path)
    
    markdown_content = []
    
    # 处理每个段落
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        if not text:
            markdown_content.append('')
            continue
        
        # 处理标题
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name.split()[-1])
            markdown_content.append('#' * level + ' ' + text)
        else:
            # 处理普通段落的格式
            formatted_text = process_paragraph_formatting(paragraph)
            markdown_content.append(formatted_text)
    
    # 处理表格
    for table in doc.tables:
        markdown_content.append('')
        markdown_content.extend(convert_table_to_markdown(table))
        markdown_content.append('')
    
    # 写入Markdown文件
    with open(output_md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(markdown_content))
    
    print(f"转换完成：{word_file_path} -> {output_md_path}")

def process_paragraph_formatting(paragraph):
    """处理段落中的文本格式"""
    result = ""
    
    for run in paragraph.runs:
        text = run.text
        
        # 处理粗体
        if run.bold:
            text = f"**{text}**"
        
        # 处理斜体
        if run.italic:
            text = f"*{text}*"
        
        # 处理下划线（转换为粗体）
        if run.underline:
            text = f"**{text}**"
        
        result += text
    
    return result

def convert_table_to_markdown(table):
    """将Word表格转换为Markdown表格"""
    markdown_table = []
    
    # 处理表头
    if table.rows:
        header_row = table.rows[0]
        header_cells = [cell.text.strip() for cell in header_row.cells]
        markdown_table.append('| ' + ' | '.join(header_cells) + ' |')
        markdown_table.append('| ' + ' | '.join(['---'] * len(header_cells)) + ' |')
        
        # 处理数据行
        for row in table.rows[1:]:
            data_cells = [cell.text.strip() for cell in row.cells]
            markdown_table.append('| ' + ' | '.join(data_cells) + ' |')
    
    return markdown_table

def batch_convert_word_files(input_directory, output_directory=None):
    """
    批量转换目录中的所有Word文档
    
    Args:
        input_directory: 包含Word文档的目录
        output_directory: 输出目录，如果为None则在原目录生成
    """
    if output_directory is None:
        output_directory = input_directory
    
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
    
    word_extensions = ['.docx', '.doc']
    converted_count = 0
    
    for filename in os.listdir(input_directory):
        file_path = os.path.join(input_directory, filename)
        
        if os.path.isfile(file_path) and any(filename.lower().endswith(ext) for ext in word_extensions):
            base_name = os.path.splitext(filename)[0]
            output_path = os.path.join(output_directory, base_name + '.md')
            
            try:
                convert_word_to_markdown(file_path, output_path)
                converted_count += 1
            except Exception as e:
                print(f"转换失败 {filename}: {e}")
    
    print(f"批量转换完成，共转换 {converted_count} 个文件")

if __name__ == "__main__":
    # 使用示例
    print("Word文档转Markdown工具")
    print("1. 转换单个文件")
    print("2. 批量转换目录中的所有Word文档")
    
    choice = input("请选择操作 (1/2): ").strip()
    
    if choice == '1':
        word_file = input("请输入Word文档路径: ").strip()
        output_file = input("请输入输出Markdown文件路径 (留空自动生成): ").strip()
        
        if not output_file:
            output_file = None
        
        convert_word_to_markdown(word_file, output_file)
    
    elif choice == '2':
        input_dir = input("请输入包含Word文档的目录路径: ").strip()
        output_dir = input("请输入输出目录路径 (留空在原目录生成): ").strip()
        
        if not output_dir:
            output_dir = None
        
        batch_convert_word_files(input_dir, output_dir)
    
    else:
        print("无效选择")